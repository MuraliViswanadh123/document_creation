"""
DOCX assembler — produces final IEEE-formatted Word document.
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def assemble_docx(sections: list, diagrams: list, profile: dict) -> bytes:
    """
    Stitch sections + diagrams into a final IEEE-formatted DOCX.
    
    Returns: DOCX file as bytes (ready for download)
    """
    doc = Document()
    _setup_styles(doc)

    # ==== TITLE PAGE ====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(profile.get("title", "Project Documentation"))
    title_run.bold = True
    title_run.font.size = Pt(20)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"\n{profile.get('domain', '')}")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    # Project metadata block
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Type: {profile.get('project_type', '').replace('_', ' ').title()}\n").italic = True
    meta.add_run(f"Primary Language: {profile.get('primary_language', '')}\n").italic = True
    if profile.get("tech_stack"):
        meta.add_run(f"Tech Stack: {', '.join(profile['tech_stack'][:8])}").italic = True

    doc.add_page_break()

    # Map diagram placement to sections (after these sections, insert relevant diagrams)
    diagram_after_section = {
        "System Architecture": [d for d in diagrams if d["name"] in ("Component Diagram", "Deployment Diagram")],
        "Methodology": [d for d in diagrams if d["name"] in ("Activity Diagram", "Sequence Diagram")],
        "Functional Requirements": [d for d in diagrams if d["name"] == "Use Case Diagram"],
        "Implementation": [d for d in diagrams if d["name"] == "Class Diagram"],
    }
    used_diagrams = set()

    # ==== SECTIONS ====
    for i, section in enumerate(sections, 1):
        # Section heading
        heading = doc.add_heading(f"{i}. {section['name']}", level=1)
        
        # Section content (parse markdown-ish content)
        _add_markdown_content(doc, section["content"])
        
        # Insert relevant diagrams after specific sections
        section_diagrams = diagram_after_section.get(section["name"], [])
        for diagram in section_diagrams:
            if diagram["name"] not in used_diagrams:
                _add_diagram(doc, diagram)
                used_diagrams.add(diagram["name"])
        
        doc.add_paragraph()  # spacing

    # ==== APPENDIX: Any remaining diagrams ====
    remaining_diagrams = [d for d in diagrams if d["name"] not in used_diagrams]
    if remaining_diagrams:
        doc.add_page_break()
        doc.add_heading("Appendix A: Additional Diagrams", level=1)
        for diagram in remaining_diagrams:
            _add_diagram(doc, diagram)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# HELPERS
# ============================================================
def _setup_styles(doc: Document):
    """Configure default styles for academic appearance."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    
    # Paragraph spacing
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _add_markdown_content(doc: Document, content: str):
    """
    Parse simple markdown content and add to document.
    Handles: ### sub-headings, **bold**, *italic*, bullet points, numbered lists.
    """
    lines = content.split("\n")
    
    for line in lines:
        line = line.rstrip()
        if not line:
            continue
        
        # Handle sub-headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            # Skip - section title already added
            continue
        # Handle bullet points
        elif line.lstrip().startswith(("- ", "* ")):
            text = line.lstrip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, text)
        # Handle numbered lists
        elif re.match(r"^\d+\.\s", line.lstrip()):
            text = re.sub(r"^\d+\.\s", "", line.lstrip())
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, text)
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)


def _add_formatted_runs(paragraph, text: str):
    """
    Add text to a paragraph, handling **bold** and *italic* markdown.
    """
    # Pattern matches **bold** or *italic*
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_diagram(doc: Document, diagram: dict):
    """Insert a diagram (image + caption) into the document."""
    if not diagram.get("image_bytes"):
        # If image rendering failed, add a placeholder note
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{diagram['name']} — image unavailable]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        return
    
    try:
        img_stream = io.BytesIO(diagram["image_bytes"])
        doc.add_picture(img_stream, width=Inches(5.5))
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"Figure: {diagram['name']}")
        cap_run.italic = True
        cap_run.font.size = Pt(10)
    except Exception as e:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"[{diagram['name']} — failed to embed: {e}]").italic = True
